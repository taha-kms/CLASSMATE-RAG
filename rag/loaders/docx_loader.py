"""
DOCX loader.

Reads paragraphs in order and joins them into a single text block.
We return a single (1, text) tuple; chunking will split further down the pipeline.
"""

from __future__ import annotations

from pathlib import Path
from typing import Tuple, List

from docx import Document
from rag.utils.text import normalize_text


def load_docx_blocks(path: str | Path) -> List[Tuple[int, str]]:
    p = Path(path)
    if not p.exists() or p.suffix.lower() != ".docx":
        raise ValueError(f"Expected an existing .docx file, got: {path}")

    doc = Document(str(p))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    # Tables (optional basic extraction)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = normalize_text("\n".join(parts))
    return [(1, text)]
